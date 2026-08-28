"""Geração de relatório semanal rico em .docx.

Cruza as 3 fontes (Proxxima, painel-ope, recorrência) e gera um documento
com resumo executivo, tendências, padrões de recorrência, análise de
produtividade por técnico, distribuição por natureza e comparativo com
o período anterior.
"""

from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from sqlalchemy import Integer, case, func, select
from sqlalchemy.orm import Session

from app.config import settings
from app.models.banco_horas_semanal import BancoHorasSemanal
from app.models.metrica_totvs import MetricaTotvs
from app.models.ocorrencia_recorrencia import OcorrenciaRecorrencia
from app.models.relatorio import Relatorio
from app.models.solicitacao_servico import SolicitacaoServico
from app.services.cruzamento import (
    buscar_banco_horas_unidade,
    buscar_infracoes_unidade,
    normalizar_unidade,
)
from app.services.totvs_client import REPORT_PONTUACAO_DIA_TECNICO, TotvsClient

_DIR_RELATORIOS = Path(settings.dir_relatorios or "relatorios")
_NATUREZAS_EXCLUIDAS = ("RECOLHIMENTO", "RECOLHIMENTO AGENDADO")


def _ensure_dir() -> Path:
    _DIR_RELATORIOS.mkdir(parents=True, exist_ok=True)
    return _DIR_RELATORIOS


# ── Helpers de formatação ──────────────────────────────────────────

def _addTitulo(doc: Document, texto: str) -> None:
    h = doc.add_heading(texto, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)


def _addSecao(doc: Document, titulo: str) -> None:
    h = doc.add_heading(titulo, level=2)
    for run in h.runs:
        run.font.size = Pt(12)


def _addSubsecao(doc: Document, titulo: str) -> None:
    h = doc.add_heading(titulo, level=3)
    for run in h.runs:
        run.font.size = Pt(11)


def _addParagrafo(doc: Document, texto: str, negrito: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrito
    run.font.size = Pt(10)


def _addParagrafos(doc: Document, linhas: list[str]) -> None:
    for texto in linhas:
        _addParagrafo(doc, texto)


def _addTabela(doc: Document, cabecalhos: list[str], linhas: list[list[str]]) -> None:
    if not linhas:
        _addParagrafo(doc, "  Nenhum registro no período.")
        return
    tabela = doc.add_table(rows=1 + len(linhas), cols=len(cabecalhos))
    tabela.style = "Light Grid Accent 1"
    for i, h in enumerate(cabecalhos):
        cell = tabela.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, linha in enumerate(linhas):
        for c, valor in enumerate(linha):
            cell = tabela.rows[r + 1].cells[c]
            cell.text = str(valor)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def _delta_str(atual: float, anterior: float) -> str:
    """Mostra variação: '+12', '-5', '=0'."""
    diff = atual - anterior
    if diff == 0:
        return "=0"
    sinal = "+" if diff > 0 else ""
    return f"{sinal}{diff:.0f}"


def _delta_pct(atual: float, anterior: float) -> str:
    if anterior == 0:
        return "N/A" if atual > 0 else "=0"
    pct = ((atual - anterior) / anterior) * 100
    sinal = "+" if pct > 0 else ""
    return f"{sinal}{pct:.0f}%"


# ── Consultas ao banco ─────────────────────────────────────────────

def _buscar_status(db: Session, unidade: str, de: date, ate: date) -> dict:
    """Backlog + produtividade + canceladas + recorrências."""
    base = (SolicitacaoServico.unidade.ilike(f"%{unidade}%"))

    abertas = db.scalar(
        select(func.count()).select_from(SolicitacaoServico).where(
            base
            & (SolicitacaoServico.status.notilike("Fechada%"))
            & (SolicitacaoServico.status.ilike("Aberta%"))
            & (SolicitacaoServico.natureza.is_not(None))
            & (SolicitacaoServico.natureza.notin_(_NATUREZAS_EXCLUIDAS))
            & (SolicitacaoServico.natureza != "")
        )
    ) or 0

    # Fechadas por status no período
    linhas = db.execute(
        select(SolicitacaoServico.status, func.count()).where(
            base
            & (SolicitacaoServico.abertura >= de)
            & (SolicitacaoServico.abertura <= ate)
            & (SolicitacaoServico.status.notilike("Aberta%"))
            & (SolicitacaoServico.status != "Cancelado")
        ).group_by(SolicitacaoServico.status)
    ).all()

    fech_prod = fech_improd = 0
    for status, count in linhas:
        if status and status.startswith("Fechada Produtiva"):
            fech_prod += count
        elif status and status.startswith("Fechada Improdutiva"):
            fech_improd += count

    canceladas = db.scalar(
        select(func.count()).select_from(SolicitacaoServico).where(
            base
            & (SolicitacaoServico.abertura >= de)
            & (SolicitacaoServico.abertura <= ate)
            & (SolicitacaoServico.status == "Cancelado")
        )
    ) or 0

    total_fechadas = db.scalar(
        select(func.count()).select_from(SolicitacaoServico).where(
            base
            & (SolicitacaoServico.abertura >= de)
            & (SolicitacaoServico.abertura <= ate)
            & (SolicitacaoServico.status.notilike("Aberta%"))
        )
    ) or 0

    recorrencias = db.scalar(
        select(func.count()).select_from(OcorrenciaRecorrencia).where(
            (OcorrenciaRecorrencia.data_abertura >= de)
            & (OcorrenciaRecorrencia.data_abertura <= ate)
            & (OcorrenciaRecorrencia.e_recorrencia.is_(True))
            & (OcorrenciaRecorrencia.unidade.ilike(f"%{unidade}%"))
        )
    ) or 0

    return {
        "abertas": abertas,
        "fech_prod": fech_prod,
        "fech_improd": fech_improd,
        "canceladas": canceladas,
        "total_fechadas": total_fechadas,
        "recorrencias": recorrencias,
    }


def _buscar_produtividade_por_tecnico(db: Session, unidade: str, de: date, ate: date, limite: int = 15) -> list[dict]:
    """Produtividade detalhada por técnico (abertas, fech_prod, fech_improd, canceladas, total)."""
    base = (SolicitacaoServico.unidade.ilike(f"%{unidade}%"))

    rows = db.execute(
        select(
            SolicitacaoServico.tecnico,
            SolicitacaoServico.status,
            func.count(),
        ).where(
            base
            & (SolicitacaoServico.abertura >= de)
            & (SolicitacaoServico.abertura <= ate)
            & (SolicitacaoServico.tecnico.isnot(None))
        ).group_by(SolicitacaoServico.tecnico, SolicitacaoServico.status)
    ).all()

    agg: dict[str, dict] = {}
    for tecnico, status, count in rows:
        if tecnico not in agg:
            agg[tecnico] = {"tecnico": tecnico, "abertas": 0, "fech_prod": 0, "fech_improd": 0, "canceladas": 0, "total": 0}
        if status and _is_aberta(status):
            agg[tecnico]["abertas"] += count
        elif status and status.startswith("Fechada Produtiva"):
            agg[tecnico]["fech_prod"] += count
        elif status and status.startswith("Fechada Improdutiva"):
            agg[tecnico]["fech_improd"] += count
        elif status and status.strip().lower() == "cancelado":
            agg[tecnico]["canceladas"] += count
        agg[tecnico]["total"] += count

    lista = sorted(agg.values(), key=lambda x: x["total"], reverse=True)[:limite]
    for item in lista:
        total = item["total"] or 1
        item["taxa_produtividade"] = round((item["fech_prod"] / total) * 100, 1)
    return lista


def _is_aberta(status: str | None) -> bool:
    if not status:
        return False
    return status.lower().startswith("aberta") or status.lower() == "aberto"


def _buscar_recorrencia_por_tecnico(db: Session, unidade: str, de: date, ate: date, limite: int = 15) -> list[dict]:
    """Total de protocolos e reaberturas por técnico."""
    rows = db.execute(
        select(
            OcorrenciaRecorrencia.tecnico,
            func.count().label("total"),
            func.sum(func.cast(OcorrenciaRecorrencia.e_recorrencia, Integer)).label("reaberturas"),
        ).where(
            OcorrenciaRecorrencia.data_abertura >= de,
            OcorrenciaRecorrencia.data_abertura <= ate,
            OcorrenciaRecorrencia.tecnico.isnot(None),
            OcorrenciaRecorrencia.unidade.ilike(f"%{unidade}%"),
        ).group_by(OcorrenciaRecorrencia.tecnico)
        .order_by(func.sum(func.cast(OcorrenciaRecorrencia.e_recorrencia, Integer)).desc())
        .limit(limite)
    ).all()

    return [
        {"tecnico": r.tecnico, "total": r.total, "reaberturas": int(r.reaberturas or 0)}
        for r in rows
    ]


def _buscar_top_he(db: Session, de: date, ate: date, limite: int = 15) -> list[dict]:
    """Top técnicos com mais HE via snapshots rankTecHE."""
    he_por_tecnico: dict[str, float] = {}
    snaps = db.scalars(
        select(BancoHorasSemanal).where(
            BancoHorasSemanal.semana_ate >= de,
            BancoHorasSemanal.semana_de <= ate,
        )
    ).all()
    for snap in snaps:
        payload = snap.payload or {}
        for item in payload.get("rankTecHE") or []:
            if isinstance(item, dict):
                nome = item.get("nome", "")
                he = float(item.get("heHoras") or 0)
                he_por_tecnico[nome] = he_por_tecnico.get(nome, 0) + he

    ordenados = sorted(he_por_tecnico.items(), key=lambda x: x[1], reverse=True)[:limite]
    return [{"tecnico": nome, "he_horas": round(he, 2)} for nome, he in ordenados]


def _buscar_naturezas(db: Session, unidade: str, de: date, ate: date, limite: int = 10) -> list[dict]:
    """Distribuição de OS por natureza no período."""
    rows = db.execute(
        select(
            SolicitacaoServico.natureza,
            func.count(),
        ).where(
            SolicitacaoServico.unidade.ilike(f"%{unidade}%")
            & (SolicitacaoServico.abertura >= de)
            & (SolicitacaoServico.abertura <= ate)
            & (SolicitacaoServico.natureza.is_not(None))
            & (SolicitacaoServico.natureza != "")
        ).group_by(SolicitacaoServico.natureza)
        .order_by(func.count().desc())
        .limit(limite)
    ).all()
    return [{"natureza": n or "(vazio)", "qtd": c} for n, c in rows]


def _buscar_distribuicao_dia_semana(db: Session, unidade: str, de: date, ate: date) -> list[dict]:
    """Distribuição de aberturas por dia da semana (0=seg, 6=dom)."""
    rows = db.execute(
        select(
            func.extract("dow", SolicitacaoServico.abertura).label("dow"),
            func.count(),
        ).where(
            SolicitacaoServico.unidade.ilike(f"%{unidade}%")
            & (SolicitacaoServico.abertura >= de)
            & (SolicitacaoServico.abertura <= ate)
        ).group_by("dow")
        .order_by("dow")
    ).all()

    dias = {0: "Seg", 1: "Ter", 2: "Qua", 3: "Qui", 4: "Sex", 5: "Sáb", 6: "Dom"}
    return [{"dia": dias.get(int(dow), str(dow)), "qtd": c} for dow, c in rows]


def _buscar_top_protocolos_recorrentes(db: Session, unidade: str, de: date, ate: date, limite: int = 10) -> list[dict]:
    """Protocolos que mais geraram recorrência (com problema de fechamento)."""
    rows = db.execute(
        select(
            OcorrenciaRecorrencia.protocolo,
            OcorrenciaRecorrencia.tecnico,
            OcorrenciaRecorrencia.problema_fechamento,
            OcorrenciaRecorrencia.dias_entre_os,
        ).where(
            (OcorrenciaRecorrencia.data_abertura >= de)
            & (OcorrenciaRecorrencia.data_abertura <= ate)
            & (OcorrenciaRecorrencia.e_recorrencia.is_(True))
            & (OcorrenciaRecorrencia.unidade.ilike(f"%{unidade}%"))
        ).order_by(OcorrenciaRecorrencia.data_abertura.desc())
        .limit(limite)
    ).all()

    return [
        {
            "protocolo": r.protocolo,
            "tecnico": r.tecnico or "?",
            "problema": (r.problema_fechamento or "")[:60],
            "dias": r.dias_entre_os,
        }
        for r in rows
    ]


def _buscar_tecnicos_com_he_e_recorrencia(db: Session, unidade: str, de: date, ate: date) -> list[dict]:
    """Técnicos que têm HE E recorrência — risco combinado."""
    # HE por técnico
    he_map: dict[str, float] = {}
    snaps = db.scalars(
        select(BancoHorasSemanal).where(
            BancoHorasSemanal.semana_ate >= de,
            BancoHorasSemanal.semana_de <= ate,
        )
    ).all()
    for snap in snaps:
        for item in (snap.payload or {}).get("rankTecHE") or []:
            if isinstance(item, dict):
                nome = item.get("nome", "")
                he_map[nome] = he_map.get(nome, 0) + float(item.get("heHoras") or 0)

    # Recorrência por técnico
    rec_map: dict[str, int] = {}
    rows = db.execute(
        select(
            OcorrenciaRecorrencia.tecnico,
            func.sum(func.cast(OcorrenciaRecorrencia.e_recorrencia, Integer)),
        ).where(
            OcorrenciaRecorrencia.data_abertura >= de,
            OcorrenciaRecorrencia.data_abertura <= ate,
            OcorrenciaRecorrencia.tecnico.isnot(None),
            OcorrenciaRecorrencia.unidade.ilike(f"%{unidade}%"),
        ).group_by(OcorrenciaRecorrencia.tecnico)
    ).all()
    for tecnico, rec in rows:
        if rec:
            rec_map[tecnico] = int(rec)

    # Cruza
    ambos = []
    for nome in set(he_map) & set(rec_map):
        ambos.append({
            "tecnico": nome,
            "he_horas": round(he_map[nome], 2),
            "reaberturas": rec_map[nome],
        })
    ambos.sort(key=lambda x: x["reaberturas"], reverse=True)
    return ambos


def _buscar_pontuacao_totvs_por_tecnico(
    db: Session, unidade: str, de: date, ate: date, limite: int = 20
) -> list[dict]:
    """Pontuação TOTVS por técnico da unidade no período.

    Lê o snapshot mais recente de ``metrica_totvs`` (report 2837323),
    parseia o xtab hierárquico e agrega por técnico.
    """
    from datetime import datetime as _dt

    reg = db.scalars(
        select(MetricaTotvs)
        .where(MetricaTotvs.report_id == REPORT_PONTUACAO_DIA_TECNICO)
        .order_by(MetricaTotvs.data_referencia.desc())
        .limit(1)
    ).first()
    if not reg or not reg.payload:
        return []

    xtab = reg.payload.get("xtab_data", reg.payload)
    dados = TotvsClient.parse_xtab_data(xtab)

    agg: dict[str, dict] = {}

    for r in dados:
        r_unidade = normalizar_unidade(r.get("unidade", ""))
        if r_unidade != unidade:
            continue
        try:
            data_ref = _dt.strptime(r["data"], "%d/%m/%Y").date()
        except (ValueError, KeyError):
            continue
        if data_ref < de or data_ref > ate:
            continue
        try:
            pontuacao = float(r.get("pontuacao", 0))
        except (ValueError, TypeError):
            pontuacao = 0.0

        tech = r.get("tecnico", "")
        if tech not in agg:
            agg[tech] = {"tecnico": tech, "total": 0.0, "dias": 0, "melhor": 0.0, "pior": 999.0}
        agg[tech]["total"] += pontuacao
        agg[tech]["dias"] += 1
        agg[tech]["melhor"] = max(agg[tech]["melhor"], pontuacao)
        agg[tech]["pior"] = min(agg[tech]["pior"], pontuacao)

    lista = []
    for item in agg.values():
        item["media"] = round(item["total"] / item["dias"], 2) if item["dias"] > 0 else 0
        item["total"] = round(item["total"], 2)
        item["melhor"] = round(item["melhor"], 2)
        item["pior"] = round(item["pior"], 2)
        lista.append(item)

    lista.sort(key=lambda x: x["media"], reverse=True)
    return lista[:limite]


# ── Gerador do relatório ───────────────────────────────────────────

def gerar_relatorio_semanal(
    db: Session,
    unidade: str,
    periodo_de: date,
    periodo_ate: date,
) -> Relatorio:
    """Gera relatório semanal rico em .docx e salva no disco + banco."""
    unidade_norm = normalizar_unidade(unidade)
    titulo = f"Relatório Semanal — {unidade_norm} ({periodo_de.strftime('%d/%m/%Y')} a {periodo_ate.strftime('%d/%m/%Y')})"

    # Período anterior (mesma duração)
    duracao = (periodo_ate - periodo_de).days
    anterior_de = periodo_de - timedelta(days=duracao + 1)
    anterior_ate = periodo_de - timedelta(days=1)

    # Coleta dados — período atual
    status = _buscar_status(db, unidade_norm, periodo_de, periodo_ate)
    bh = buscar_banco_horas_unidade(db, unidade_norm, periodo_de, periodo_ate)
    infracoes = buscar_infracoes_unidade(db, unidade_norm, periodo_de, periodo_ate)
    prod_tec = _buscar_produtividade_por_tecnico(db, unidade_norm, periodo_de, periodo_ate)
    rec_tec = _buscar_recorrencia_por_tecnico(db, unidade_norm, periodo_de, periodo_ate)
    top_he = _buscar_top_he(db, periodo_de, periodo_ate)
    naturezas = _buscar_naturezas(db, unidade_norm, periodo_de, periodo_ate)
    dist_dia = _buscar_distribuicao_dia_semana(db, unidade_norm, periodo_de, periodo_ate)
    prot_rec = _buscar_top_protocolos_recorrentes(db, unidade_norm, periodo_de, periodo_ate)
    risco = _buscar_tecnicos_com_he_e_recorrencia(db, unidade_norm, periodo_de, periodo_ate)

    # Coleta dados — período anterior (para tendências)
    status_ant = _buscar_status(db, unidade_norm, anterior_de, anterior_ate)
    bh_ant = buscar_banco_horas_unidade(db, unidade_norm, anterior_de, anterior_ate)

    # ── Monta documento ──────────────────────────────────────────
    doc = Document()
    _addTitulo(doc, titulo)

    # 1. Resumo Executivo com Tendências
    _addSecao(doc, "1. Resumo Executivo")
    _addParagrafos(doc, [
        f"Unidade: {unidade_norm}",
        f"Período analisado: {periodo_de.strftime('%d/%m/%Y')} a {periodo_ate.strftime('%d/%m/%Y')}",
        f"Período de comparação: {anterior_de.strftime('%d/%m/%Y')} a {anterior_ate.strftime('%d/%m/%Y')}",
    ])

    _addSubsecao(doc, "Indicadores-chave")
    _addTabela(doc,
        ["Indicador", "Atual", "Anterior", "Variação"],
        [
            ["Backlog (abertas)", str(status["abertas"]), str(status_ant["abertas"]),
             _delta_str(status["abertas"], status_ant["abertas"])],
            ["Fechadas produtivas", str(status["fech_prod"]), str(status_ant["fech_prod"]),
             _delta_str(status["fech_prod"], status_ant["fech_prod"])],
            ["Fechadas improdutivas", str(status["fech_improd"]), str(status_ant["fech_improd"]),
             _delta_str(status["fech_improd"], status_ant["fech_improd"])],
            ["Canceladas", str(status["canceladas"]), str(status_ant["canceladas"]),
             _delta_str(status["canceladas"], status_ant["canceladas"])],
            ["Horas Extras", f"{bh['he_horas']:.1f}h", f"{bh_ant['he_horas']:.1f}h",
             _delta_str(bh["he_horas"], bh_ant["he_horas"])],
            ["Infrações", str(infracoes), "-", "-"],
            ["Recorrências", str(status["recorrencias"]), str(status_ant["recorrencias"]),
             _delta_str(status["recorrencias"], status_ant["recorrencias"])],
        ]
    )

    # Taxa de produtividade
    total_fechadas = status["fech_prod"] + status["fech_improd"]
    taxa = round((status["fech_prod"] / total_fechadas * 100), 1) if total_fechadas > 0 else 0
    total_ant = status_ant["fech_prod"] + status_ant["fech_improd"]
    taxa_ant = round((status_ant["fech_prod"] / total_ant * 100), 1) if total_ant > 0 else 0
    _addParagrafo(doc, f"Taxa de produtividade: {taxa}% (anterior: {taxa_ant}%)", negrito=True)

    # 2. Análise de Tendências
    _addSecao(doc, "2. Análise de Tendências")
    delta_abertas = status["abertas"] - status_ant["abertas"]
    delta_prod = status["fech_prod"] - status_ant["fech_prod"]
    delta_he = bh["he_horas"] - bh_ant["he_horas"]
    delta_rec = status["recorrencias"] - status_ant["recorrencias"]

    insights = []
    if delta_abertas > 20:
        insights.append(f"ATENÇÃO: Backlog cresceu {delta_abertas} OS — acumulação acima do normal.")
    elif delta_abertas < -20:
        insights.append(f"POSITIVO: Backlog reduziu {abs(delta_abertas)} OS — ritmo de fechamento acima da abertura.")
    else:
        insights.append(f"Backlog estável (variação de {delta_abertas}).")

    if delta_prod > 10:
        insights.append(f"Produtividade subiu {delta_prod} fechadas — equipe mais produtiva que na semana anterior.")
    elif delta_prod < -10:
        insights.append(f"Produtividade caiu {abs(delta_prod)} fechadas — possível gargalo ou falta de equipe.")

    if delta_he > 10:
        insights.append(f"HE cresceu {delta_he:.1f}h — carga de trabalho acima do planejado.")
    elif delta_he < -10:
        insights.append(f"HE reduziu {abs(delta_he):.1f}h — equipe dentro da capacidade normal.")

    if delta_rec > 5:
        insights.append(f"Recorrências subiram {delta_rec} — qualidade dos fechamentos caiu.")
    elif delta_rec < -5:
        insights.append(f"Recorrências reduziram {abs(delta_rec)} — melhoria na qualidade.")

    if not insights:
        insights.append("Sem alterações significativas em relação à semana anterior.")

    _addParagrafos(doc, insights)

    # 3. Análise de Produtividade por Técnico
    _addSecao(doc, "3. Produtividade por Técnico")
    _addParagrafo(doc, "Top técnicos por volume total de OS no período:")

    if prod_tec:
        _addTabela(doc,
            ["Técnico", "Abertas", "Produtivas", "Improdutivas", "Canceladas", "Total", "Taxa Prod."],
            [
                [
                    item["tecnico"],
                    str(item["abertas"]),
                    str(item["fech_prod"]),
                    str(item["fech_improd"]),
                    str(item["canceladas"]),
                    str(item["total"]),
                    f"{item['taxa_produtividade']}%",
                ]
                for item in prod_tec
            ]
        )

        # Insights de produtividade
        taxas = [(item["tecnico"], item["taxa_produtividade"]) for item in prod_tec if item["total"] >= 3]
        if taxas:
            taxas.sort(key=lambda x: x[1])
            pior = taxas[0]
            melhor = taxas[-1]
            _addParagrafo(doc, f"Maior taxa: {melhor[0]} ({melhor[1]}%)", negrito=True)
            _addParagrafo(doc, f"Menor taxa: {pior[0]} ({pior[1]}%) — considere acompanhamento.", negrito=True)
    else:
        _addParagrafo(doc, "  Nenhum dado de produtividade no período.")

    # 4. Recorrência por Técnico
    _addSecao(doc, "4. Recorrência por Técnico")
    if rec_tec:
        _addTabela(doc,
            ["Técnico", "Protocolos", "Reaberturas", "Taxa Recorrência"],
            [
                [
                    r["tecnico"],
                    str(r["total"]),
                    str(r["reaberturas"]),
                    f"{round(r['reaberturas'] / r['total'] * 100, 1) if r['total'] > 0 else 0}%",
                ]
                for r in rec_tec
            ]
        )

        total_rec_geral = sum(r["reaberturas"] for r in rec_tec)
        if total_rec_geral > 0:
            top3 = rec_tec[:3]
            top3_total = sum(r["reaberturas"] for r in top3)
            pct_top3 = round((top3_total / total_rec_geral) * 100, 0)
            _addParagrafo(doc, f"Top 3 técnicos concentram {pct_top3:.0f}% das reaberturas ({top3_total}/{total_rec_geral}).", negrito=True)
    else:
        _addParagrafo(doc, "  Nenhuma recorrência registrada no período.")

    # 5. Horas Extras
    _addSecao(doc, "5. Horas Extras por Técnico")
    if top_he:
        _addTabela(doc,
            ["Técnico", "HE (h)"],
            [[r["tecnico"], f"{r['he_horas']:.2f}"] for r in top_he]
        )
        total_he_geral = sum(r["he_horas"] for r in top_he)
        _addParagrafo(doc, f"Total HE top 10: {total_he_geral:.1f}h")
    else:
        _addParagrafo(doc, "  Nenhum registro de HE no período.")

    # 6. Distribuição por Natureza
    _addSecao(doc, "6. Distribuição por Natureza")
    if naturezas:
        total_nat = sum(n["qtd"] for n in naturezas)
        _addTabela(doc,
            ["Natureza", "Qtd", "% do Total"],
            [
                [n["natureza"], str(n["qtd"]), f"{round(n['qtd'] / total_nat * 100, 1)}%"]
                for n in naturezas
            ]
        )
    else:
        _addParagrafo(doc, "  Nenhum dado de natureza no período.")

    # 7. Padrões Temporais
    _addSecao(doc, "7. Distribuição por Dia da Semana")
    if dist_dia:
        _addTabela(doc,
            ["Dia", "Aberturas"],
            [[d["dia"], str(d["qtd"])] for d in dist_dia]
        )
        max_dia = max(dist_dia, key=lambda x: x["qtd"])
        min_dia = min(dist_dia, key=lambda x: x["qtd"])
        _addParagrafo(doc, f"Dia com mais aberturas: {max_dia['dia']} ({max_dia['qtd']} OS)", negrito=True)
        _addParagrafo(doc, f"Dia com menos aberturas: {min_dia['dia']} ({min_dia['qtd']} OS)", negrito=True)
    else:
        _addParagrafo(doc, "  Nenhum dado temporal no período.")

    # 8. Risco Combinado (HE + Recorrência)
    _addSecao(doc, "8. Técnicos com Risco Combinado (HE + Recorrência)")
    _addParagrafo(doc, "Técnicos que acumulam horas extras E recorrências — maior atenção.")
    if risco:
        _addTabela(doc,
            ["Técnico", "HE (h)", "Reaberturas"],
            [[r["tecnico"], f"{r['he_horas']:.1f}", str(r["reaberturas"])] for r in risco]
        )
    else:
        _addParagrafo(doc, "  Nenhum técnico com risco combinado no período.")

    # 9. Protocolos com Recorrência (detalhe)
    _addSecao(doc, "9. Protocolos com Recorrência (detalhe)")
    if prot_rec:
        _addTabela(doc,
            ["Protocolo", "Técnico", "Problema do Fechamento", "Dias entre OS"],
            [
                [
                    p["protocolo"],
                    p["tecnico"],
                    p["problema"],
                    str(p["dias"]) if p["dias"] is not None else "?",
                ]
                for p in prot_rec
            ]
        )
    else:
        _addParagrafo(doc, "  Nenhum protocolo recorrente no período.")

    # 10. Pontuação TOTVS (GoodData)
    pont_totvs = _buscar_pontuacao_totvs_por_tecnico(db, unidade_norm, periodo_de, periodo_ate)
    _addSecao(doc, "10. Pontuação TOTVS por Técnico")
    _addParagrafo(doc, "Média de pontuação diária (GoodData) por técnico no período.")
    if pont_totvs:
        _addTabela(doc,
            ["Técnico", "Média", "Total", "Dias", "Melhor", "Pior"],
            [
                [
                    item["tecnico"],
                    f"{item['media']:.2f}",
                    f"{item['total']:.2f}",
                    str(item["dias"]),
                    f"{item['melhor']:.2f}",
                    f"{item['pior']:.2f}",
                ]
                for item in pont_totvs
            ]
        )
        media_geral = sum(item["media"] for item in pont_totvs) / len(pont_totvs)
        _addParagrafo(doc, f"Média geral da unidade: {media_geral:.2f}", negrito=True)

        acima = [item for item in pont_totvs if item["media"] >= 7.0]
        abaixo = [item for item in pont_totvs if item["media"] < 7.0]
        if acima:
            _addParagrafo(doc, f"Técnicos acima da meta (≥7.0): {len(acima)}/{len(pont_totvs)}")
        if abaixo:
            nomes = ", ".join(item["tecnico"] for item in abaixo[:5])
            _addParagrafo(doc, f"Técnicos abaixo da meta (<7.0): {len(abaixo)}/{len(pont_totvs)} — {nomes}", negrito=True)
    else:
        _addParagrafo(doc, "  Nenhum dado de pontuação TOTVS no período.")

    # 11. Observações e Fontes
    _addSecao(doc, "11. Observações")
    _addParagrafos(doc, [
        "Relatório gerado automaticamente pelo Agente OPE.",
        "Fontes: Proxxima (solicitações), painel-ope (HE/infrações), Excel de recorrência, TOTVS Analytics (pontuação).",
        "Período de comparação: mesma duração imediatamente anterior.",
        "Para dúvidas ou aprofundamento, consulte o agente operacoes no opencode.",
        f"Gerado em: {date.today().strftime('%d/%m/%Y')}",
    ])

    # Salva em disco
    _dir = _ensure_dir()
    nome_arquivo = f"relatorio_{unidade_norm.replace(' ', '_')}_{periodo_de}_{periodo_ate}.docx"
    caminho = _dir / nome_arquivo
    doc.save(str(caminho))

    # Registra no banco
    reg = Relatorio(
        titulo=titulo,
        unidade=unidade_norm,
        periodo_de=periodo_de.isoformat(),
        periodo_ate=periodo_ate.isoformat(),
        nome_arquivo=nome_arquivo,
        caminho=str(caminho),
    )
    db.add(reg)
    db.commit()
    db.refresh(reg)

    return reg
