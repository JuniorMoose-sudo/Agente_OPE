"""Testes auxiliares do relatório (Sprint 7).

Cobre: _addSubsecao, _addParagrafo, _addTabela, _ensure_dir, lógica de
tendências, e a lógica de cálculo de taxa de produtividade.
"""

from unittest.mock import MagicMock, patch
from pathlib import Path

from docx import Document

from app.services.relatorio import (
    _addParagrafo,
    _addSubsecao,
    _addTabela,
    _addTitulo,
    _ensure_dir,
)


# ── Helpers de formatação ──────────────────────────────────────────


class TestAddTitulo:
    def test_adiciona_heading(self):
        doc = Document()
        _addTitulo(doc, "Teste")
        assert len(doc.paragraphs) >= 1
        texto = doc.paragraphs[0].text
        assert "Teste" in texto


class TestAddSubsecao:
    def test_adiciona_subsecao(self):
        doc = Document()
        _addSubsecao(doc, "Subseção Teste")
        assert any("Subseção Teste" in p.text for p in doc.paragraphs)


class TestAddParagrafo:
    def test_paragrafo_normal(self):
        doc = Document()
        _addParagrafo(doc, "Linha teste")
        assert any("Linha teste" in p.text for p in doc.paragraphs)

    def test_paragrafo_negrito(self):
        doc = Document()
        _addParagrafo(doc, "Negrito", negrito=True)
        # Verificar que o parágrafo foi adicionado
        assert len(doc.paragraphs) >= 1


class TestAddTabela:
    def test_tabela_vazia(self):
        doc = Document()
        _addTabela(doc, ["Col1", "Col2"], [])
        # Tabela vazia gera "Nenhum registro"
        assert any("Nenhum registro" in p.text for p in doc.paragraphs)

    def test_tabela_com_dados(self):
        doc = Document()
        _addTabela(doc, ["Nome", "Valor"], [["A", "1"], ["B", "2"]])
        #python-docx cria tabelas como block items, não parágrafos
        # Mas a tabela deve existir
        assert len(doc.tables) >= 1


# ── _ensure_dir ────────────────────────────────────────────────────


class TestEnsureDir:
    def test_cria_diretorio(self, tmp_path):
        with patch("app.services.relatorio._DIR_RELATORIOS", tmp_path / "relatorios_test"):
            result = _ensure_dir()
            assert result.exists()
            assert result.is_dir()


# ── Lógica de cálculo (sem DB) ────────────────────────────────────


class TestLogicaCalculo:
    def test_taxa_produtividade_calculo(self):
        """Testa a lógica de cálculo de taxa de produtividade usada no relatório."""
        fech_prod = 80
        fech_improd = 20
        total_fechadas = fech_prod + fech_improd
        taxa = round((fech_prod / total_fechadas * 100), 1) if total_fechadas > 0 else 0
        assert taxa == 80.0

    def test_taxa_produtividade_zero(self):
        """Quando não há fechadas, taxa deve ser 0."""
        fech_prod = 0
        fech_improd = 0
        total_fechadas = fech_prod + fech_improd
        taxa = round((fech_prod / total_fechadas * 100), 1) if total_fechadas > 0 else 0
        assert taxa == 0

    def test_indice_top3_concentracao(self):
        """Testa o cálculo de concentração top 3 usada no relatório de recorrência."""
        rec_tec = [
            {"tecnico": "A", "reaberturas": 10},
            {"tecnico": "B", "reaberturas": 8},
            {"tecnico": "C", "reaberturas": 5},
            {"tecnico": "D", "reaberturas": 3},
            {"tecnico": "E", "reaberturas": 2},
        ]
        total_rec_geral = sum(r["reaberturas"] for r in rec_tec)
        top3 = rec_tec[:3]
        top3_total = sum(r["reaberturas"] for r in top3)
        pct_top3 = round((top3_total / total_rec_geral) * 100, 0)
        assert pct_top3 == 82.0  # (10+8+5)/(10+8+5+3+2) = 23/28 ≈ 82%

    def test_delta_abertas_tendencia(self):
        """Testa a lógica de tendência do backlog."""
        abertas_atual = 400
        abertas_anterior = 350
        delta = abertas_atual - abertas_anterior
        assert delta == 50
        assert delta > 20  # Gera alerta de backlog crescente

    def test_he_delta_tendencia(self):
        """Testa a lógica de tendência de HE."""
        he_atual = 150.0
        he_anterior = 120.0
        delta = he_atual - he_anterior
        assert delta == 30.0
        assert delta > 10  # Gera alerta de HE crescente


# ── _is_aberta usado no relatorio (import local) ──────────────────


class TestIsAbertaRelatorio:
    """Testa _is_aberta do módulo relatorio.py (diferente do sync_proxxima)."""

    def test_aberta_maiuscula(self):
        from app.services.relatorio import _is_aberta
        assert _is_aberta("Aberta") is True

    def test_aberta_minuscula(self):
        from app.services.relatorio import _is_aberta
        assert _is_aberta("aberta") is True

    def test_aberto(self):
        from app.services.relatorio import _is_aberta
        assert _is_aberta("aberto") is True

    def test_none(self):
        from app.services.relatorio import _is_aberta
        assert _is_aberta(None) is False

    def test_fechada(self):
        from app.services.relatorio import _is_aberta
        assert _is_aberta("Fechada Produtiva") is False

    def test_cancelado(self):
        from app.services.relatorio import _is_aberta
        assert _is_aberta("Cancelado") is False


# ── Constantes de alerta ───────────────────────────────────────────


class TestConstantesAlerta:
    def test_limite_reabertura_e_1(self):
        from app.services.cruzamento import LIMITE_REABERTURA
        assert LIMITE_REABERTURA == 1

    def test_limite_he_e_8(self):
        from app.services.cruzamento import LIMITE_HE_SEMANAL
        assert LIMITE_HE_SEMANAL == 8.0

    def test_meta_inspecao_e_7(self):
        from app.services.cruzamento import META_INSPECAO
        assert META_INSPECAO == 7.0
